import os
from dotenv import load_dotenv

# Load environment variables (for local development)
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")

import streamlit as st
from langchain_groq import ChatGroq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json

st.set_page_config(page_title="AI Resume Builder", page_icon="📄", layout="wide")

st.markdown("""
<style>
    .main { background-color: #0e1117; }
    .title { text-align: center; font-size: 42px; font-weight: bold; color: #4CAF50; }
    .subtitle { text-align: center; color: #888; margin-bottom: 10px; }
    .payment-box { background: #1e2130; border-radius: 10px; padding: 20px; border: 2px solid #4CAF50; text-align: center; margin-bottom: 10px; }
    .code-box { background: #1e2130; border-radius: 10px; padding: 20px; border: 2px solid #2196F3; text-align: center; }
    .step { background: #0e1117; border-radius: 8px; padding: 10px 15px; margin: 5px 0; border-left: 4px solid #4CAF50; text-align: left; }
</style>
""", unsafe_allow_html=True)

# ─── Premium Resume Builder ─────────────────────────────────────
def add_horizontal_line(paragraph, color="2E75B6", size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def create_premium_resume(data):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    PRIMARY = RGBColor(0x1A, 0x3A, 0x5C)
    ACCENT = RGBColor(0x2E, 0x75, 0xB6)
    DARK_GRAY = RGBColor(0x40, 0x40, 0x40)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LIGHT_BLUE = RGBColor(0xAD, 0xD8, 0xE6)

    # ── Header Table ──
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    set_cell_background(left_cell, '1A3A5C')
    set_cell_background(right_cell, '2E75B6')
    remove_cell_borders(left_cell)
    remove_cell_borders(right_cell)

    # Name
    name_para = left_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_before = Pt(10)
    name_para.paragraph_format.space_after = Pt(2)
    name_para.paragraph_format.left_indent = Inches(0.15)
    name_run = name_para.add_run(data.get('name', '').upper())
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = WHITE

    title_para = left_cell.add_paragraph()
    title_para.paragraph_format.left_indent = Inches(0.15)
    title_para.paragraph_format.space_after = Pt(10)
    title_run = title_para.add_run(data.get('job_title', ''))
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = LIGHT_BLUE
    title_run.font.italic = True

    contact_info = []
    if data.get('email'): contact_info.append(f"Email: {data['email']}")
    if data.get('phone'): contact_info.append(f"Phone: {data['phone']}")
    if data.get('location'): contact_info.append(f"Location: {data['location']}")
    if data.get('linkedin'): contact_info.append(f"LinkedIn: {data['linkedin']}")
    if data.get('github'): contact_info.append(f"GitHub: {data['github']}")

    first = True
    for info in contact_info:
        p = right_cell.paragraphs[0] if first else right_cell.add_paragraph()
        first = False
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_indent = Inches(0.1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(info)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = ACCENT
        add_horizontal_line(p)

    def two_col_line(left, right, left_bold=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '8640')
        tabs.append(tab)
        pPr.append(tabs)
        lr = p.add_run(left)
        lr.font.name = 'Calibri'
        lr.font.size = Pt(10)
        lr.font.bold = left_bold
        lr.font.color.rgb = PRIMARY
        p.add_run('\t')
        rr = p.add_run(right)
        rr.font.name = 'Calibri'
        rr.font.size = Pt(9)
        rr.font.italic = True
        rr.font.color.rgb = ACCENT

    def bullet_line(text, indent=0.2):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"• {text}")
        run.font.name = 'Calibri'
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_GRAY

    # ── Summary ──
    if data.get('summary'):
        section_header('Professional Summary')
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(data['summary'])
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY

    # ── Experience ──
    experiences = [e for e in data.get('experience', []) if e.get('title')]
    if experiences:
        section_header('Professional Experience')
        for exp in experiences:
            two_col_line(f"{exp.get('title','')}  |  {exp.get('company','')}", exp.get('duration',''))
            if exp.get('description'):
                for line in exp['description'].split('\n'):
                    if line.strip():
                        bullet_line(line.strip())

    # ── Education ──
    education = [e for e in data.get('education', []) if e.get('degree')]
    if education:
        section_header('Education')
        for edu in education:
            gpa_str = f"| GPA: {edu['gpa']}" if edu.get('gpa') else ''
            two_col_line(
                f"{edu.get('degree','')}  |  {edu.get('university','')}",
                f"{edu.get('year','')} {gpa_str}"
            )

    # ── Skills ──
    if data.get('technical_skills'):
        section_header('Skills')
        skills_table = doc.add_table(rows=1, cols=2)
        skills_table.style = 'Table Grid'
        lc = skills_table.rows[0].cells[0]
        rc = skills_table.rows[0].cells[1]
        remove_cell_borders(lc)
        remove_cell_borders(rc)

        p = lc.paragraphs[0]
        r = p.add_run('Technical Skills')
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = PRIMARY
        for skill in data['technical_skills'].split(','):
            sp = lc.add_paragraph()
            sp.paragraph_format.space_before = Pt(1)
            run = sp.add_run(f"  ▸  {skill.strip()}")
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY

        p2 = rc.paragraphs[0]
        r2 = p2.add_run('Other Skills')
        r2.font.bold = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = PRIMARY
        if data.get('soft_skills'):
            sp = rc.add_paragraph()
            run = sp.add_run(data['soft_skills'])
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
        if data.get('languages'):
            sp = rc.add_paragraph()
            run = sp.add_run(f"Languages: {data['languages']}")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = DARK_GRAY

    # ── Projects ──
    if data.get('projects'):
        section_header('Projects')
        for project in data['projects'].split('\n'):
            if project.strip():
                parts = project.split('-', 1)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(2)
                if len(parts) == 2:
                    tr = p.add_run(f"▸ {parts[0].strip()}")
                    tr.font.bold = True
                    tr.font.size = Pt(10)
                    tr.font.color.rgb = PRIMARY
                    dr = p.add_run(f" — {parts[1].strip()}")
                    dr.font.size = Pt(9)
                    dr.font.color.rgb = DARK_GRAY
                else:
                    r = p.add_run(f"▸ {project.strip()}")
                    r.font.size = Pt(10)
                    r.font.color.rgb = PRIMARY

    # ── Certifications ──
    if data.get('certifications'):
        section_header('Certifications')
        for cert in data['certifications'].split('\n'):
            if cert.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                run = p.add_run(f"✓  {cert.strip()}")
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_GRAY

    # ── Achievements ──
    if data.get('achievements'):
        section_header('Achievements & Awards')
        for ach in data['achievements'].split('\n'):
            if ach.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                run = p.add_run(f"🏆  {ach.strip()}")
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_GRAY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ─── Code Functions ─────────────────────────────────────────────
def load_codes():
    if os.path.exists("codes.json"):
        with open("codes.json", "r") as f:
            return json.load(f)
    return {}

def save_codes(codes):
    with open("codes.json", "w") as f:
        json.dump(codes, f, indent=2)

def validate_code(code):
    codes = load_codes()
    code = code.strip().upper()
    if code in codes:
        return (True, "valid") if not codes[code]["used"] else (False, "used")
    return False, "invalid"

def expire_code(code):
    codes = load_codes()
    code = code.strip().upper()
    if code in codes:
        codes[code]["used"] = True
        save_codes(codes)

# ─── Session State ──────────────────────────────────────────────
for key, val in {"resume_text": "", "resume_generated": False, "full_name": "", "resume_data": {}}.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ─── UI ─────────────────────────────────────────────────────────
st.markdown('<div class="title">🤖 AI Resume Builder</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Professional, job-specific resumes powered by AI</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">💰 Only <b>Rs 300</b> per resume</div>', unsafe_allow_html=True)
st.divider()

if st.sidebar.checkbox("🔐 Admin Panel"):
    admin_pass = st.sidebar.text_input("Password", type="password")
    if admin_pass == "shahram2004admin":
        codes = load_codes()
        total = len(codes)
        used = sum(1 for c in codes.values() if c["used"])
        st.sidebar.success(f"Total: {total} | Used: {used} | Available: {total-used}")
        if st.sidebar.button("Show Used Codes"):
            st.sidebar.write([c for c, v in codes.items() if v["used"]])
        mc = st.sidebar.text_input("Manually expire code")
        if st.sidebar.button("Expire"):
            expire_code(mc)
            st.sidebar.success("Done!")

col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("📋 Fill Your Details")
    tab1, tab2, tab3, tab4 = st.tabs(["👤 Personal", "🎓 Education", "💼 Experience", "🛠️ Skills"])

    with tab1:
        full_name = st.text_input("Full Name *", placeholder="Muhammad Shahram")
        job_title = st.text_input("Job Title Applying For *", placeholder="Machine Learning Engineer")
        email = st.text_input("Email *", placeholder="shahram@gmail.com")
        phone = st.text_input("Phone", placeholder="+92 317 0222193")
        location = st.text_input("Location", placeholder="Islamabad, Pakistan")
        linkedin = st.text_input("LinkedIn (optional)", placeholder="linkedin.com/in/shahram")
        github = st.text_input("GitHub (optional)", placeholder="github.com/shahram2004")
        summary_input = st.text_area("About Yourself", placeholder="I am a passionate engineer...")

    with tab2:
        degree1 = st.text_input("Degree 1", placeholder="BS Engineering Sciences")
        university1 = st.text_input("University 1", placeholder="NUST, Islamabad")
        year1 = st.text_input("Year 1", placeholder="2022-2026")
        gpa1 = st.text_input("GPA 1", placeholder="3.8/4.0")
        st.divider()
        degree2 = st.text_input("Degree 2 (optional)", placeholder="FSc Pre-Engineering")
        university2 = st.text_input("College 2", placeholder="Punjab College")
        year2 = st.text_input("Year 2", placeholder="2020-2022")

    with tab3:
        exp1_title = st.text_input("Job Title 1", placeholder="ML Engineer Intern")
        exp1_company = st.text_input("Company 1", placeholder="Tech Startup, Lahore")
        exp1_duration = st.text_input("Duration 1", placeholder="June 2024 - August 2024")
        exp1_desc = st.text_area("Description 1", placeholder="Worked on NLP models...")
        st.divider()
        exp2_title = st.text_input("Job Title 2 (optional)", placeholder="Freelance Developer")
        exp2_company = st.text_input("Company 2", placeholder="Fiverr")
        exp2_duration = st.text_input("Duration 2", placeholder="2023 - Present")
        exp2_desc = st.text_area("Description 2", placeholder="Built AI tools for clients...")

    with tab4:
        technical_skills = st.text_area("Technical Skills *", placeholder="Python, MATLAB, TensorFlow, LangChain, Streamlit, FPGA")
        soft_skills = st.text_area("Soft Skills", placeholder="Leadership, Communication, Problem Solving")
        languages = st.text_input("Languages", placeholder="Urdu (Native), English (Fluent)")
        certifications = st.text_area("Certifications (optional)", placeholder="Machine Learning - Coursera, 2024")
        projects = st.text_area("Projects *", placeholder="PDF AI Chatbot - Python, LangChain, Groq\nSignal Processor - MATLAB")
        achievements = st.text_area("Achievements (optional)", placeholder="Dean's List 2023, Hackathon Winner")

with col2:
    st.subheader("⚙️ Settings")
    tone = st.selectbox("Tone", ["Professional", "Creative", "Academic", "Technical"])
    resume_length = st.selectbox("Length", ["One Page", "Two Pages", "Detailed"])
    highlight = st.text_area("Paste Job Description Here", placeholder="Paste the job description for a perfectly tailored resume...")
    st.divider()

    st.markdown("""
    <div class="payment-box">
        <h3>💳 How to Get Your Resume</h3>
        <div class="step">1️⃣ Fill your details on the left</div>
        <div class="step">2️⃣ Send <b>Rs 300</b> to:<br><br>
            📱 EasyPaisa: <b>0317-0222193</b><br>
            📱 NayaPay: <b>0317-0222193</b><br>
            👤 Name: <b>Shahram Khan</b></div>
        <div class="step">3️⃣ Send screenshot on WhatsApp:<br>📞 <b>0317-0222193</b></div>
        <div class="step">4️⃣ Receive unlock code on WhatsApp</div>
        <div class="step">5️⃣ Enter code → Resume generates & downloads!</div>
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    st.markdown('<div class="code-box">', unsafe_allow_html=True)
    st.markdown("### 🔑 Enter Unlock Code")
    entered_code = st.text_input("", placeholder="e.g. ABC1234XYZ", label_visibility="collapsed")

    if st.button("🚀 Verify Code & Generate Resume", use_container_width=True, type="primary"):
        if not full_name or not job_title or not email or not technical_skills:
            st.error("⚠️ Please fill required fields (Name, Job Title, Email, Skills)")
        elif not entered_code:
            st.warning("Please enter your unlock code")
        else:
            is_valid, status = validate_code(entered_code)
            if status == "used":
                st.error("❌ This code has already been used.")
            elif not is_valid:
                st.error("❌ Invalid code. Please check and try again.")
            else:
                st.session_state.full_name = full_name
                st.session_state.resume_data = {
                    'name': full_name, 'job_title': job_title,
                    'email': email, 'phone': phone, 'location': location,
                    'linkedin': linkedin, 'github': github,
                    'summary': summary_input,
                    'experience': [
                        {'title': exp1_title, 'company': exp1_company, 'duration': exp1_duration, 'description': exp1_desc},
                        {'title': exp2_title, 'company': exp2_company, 'duration': exp2_duration, 'description': exp2_desc},
                    ],
                    'education': [
                        {'degree': degree1, 'university': university1, 'year': year1, 'gpa': gpa1},
                        {'degree': degree2, 'university': university2, 'year': year2, 'gpa': ''},
                    ],
                    'technical_skills': technical_skills, 'soft_skills': soft_skills,
                    'languages': languages, 'certifications': certifications,
                    'projects': projects, 'achievements': achievements,
                }

                with st.spinner("🤖 Generating your perfect resume..."):
                    llm = ChatGroq(model="llama-3.3-70b-versatile", api_key=GROQ_API_KEY)
                    prompt = f"""You are an expert resume writer. Create a professional ATS-optimized resume summary and bullet points only.

For this candidate applying for {job_title}:
- Name: {full_name}
- Skills: {technical_skills}
- Experience: {exp1_title} at {exp1_company}: {exp1_desc}
- Projects: {projects}
- Job Description to tailor for: {highlight}

Write ONLY:
1. A powerful 3-sentence professional summary
2. 3-4 improved bullet points for each experience (action verb + achievement)

Keep it concise and impactful. Tone: {tone}"""

                    response = llm.invoke(prompt)
                    st.session_state.resume_text = response.content
                    st.session_state.resume_generated = True
                    expire_code(entered_code.strip().upper())
                    st.success("✅ Done! Scroll down to download.")

    st.markdown('</div>', unsafe_allow_html=True)

if st.session_state.resume_generated:
    st.divider()
    st.subheader("✅ Your Premium Resume is Ready!")

    col_a, col_b = st.columns(2)
    with col_a:
        word_file = create_premium_resume(st.session_state.resume_data)
        st.download_button(
            "📥 Download Premium Word (.docx)",
            word_file,
            file_name=f"{st.session_state.full_name}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )
    with col_b:
        st.download_button(
            "📄 Download Text (.txt)",
            st.session_state.resume_text,
            file_name=f"{st.session_state.full_name}_Resume.txt",
            use_container_width=True
        )

    with st.expander("👁️ Preview AI Content"):
        st.markdown(st.session_state.resume_text)

    st.warning("⚠️ Code used and expired. Need another? WhatsApp: 0317-0222193")